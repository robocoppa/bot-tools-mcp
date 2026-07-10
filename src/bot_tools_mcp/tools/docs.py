"""Docs & sheets tools — Nextcloud files, scoped per bot.

Each tool reads/writes files under the calling bot's own Nextcloud root
(`/dav/files/<bot>/…`), authenticated as the bot. Spreadsheet ops use `openpyxl`,
document ops use `python-docx`; the file is round-tripped in memory (download
bytes → edit → upload bytes). `create_share_link` returns a public link — the one
and only place the public Nextcloud URL is used.

The `path` on every tool is guarded by `safe_path` (in the transport) against
absolute paths and `..` traversal.
"""

from __future__ import annotations

import asyncio
import io

from fastmcp import Context, FastMCP
from fastmcp.exceptions import ToolError

from bot_tools_mcp import nextcloud_client as nc
from bot_tools_mcp._util import to_tool_error
from bot_tools_mcp.identity import Identity
from bot_tools_mcp.server import bot_creds


def register(mcp: FastMCP, identity: Identity, cfg: nc.NextcloudConfig | None = None):
    """Attach the docs/sheets tools; return them as a dict for unit testing."""
    config = cfg or nc.NextcloudConfig(
        internal_url=identity.nextcloud_url(),
        public_url=identity.nextcloud_public_url(),
    )

    def _creds(ctx: Context) -> tuple[str, str]:
        return bot_creds(ctx, identity.nextcloud_password)

    async def _run(fn, *args, **kwargs):
        """Run a sync Nextcloud transport op in a thread, surfacing failures loud."""
        try:
            return await asyncio.to_thread(fn, config, *args, **kwargs)
        except Exception as exc:
            raise to_tool_error(exc) from exc

    async def _get(bot, pw, path) -> bytes:
        return await _run(nc.get_file, bot, pw, path)

    async def _put(bot, pw, path, data) -> None:
        await _run(nc.put_file, bot, pw, path, data)

    # --- spreadsheets (openpyxl) ---

    @mcp.tool
    async def sheet_create(ctx: Context, path: str, sheets: list[str] | None = None) -> str:
        """Create a new empty `.xlsx` at `path` with the given sheet names
        (default one sheet named 'Sheet1')."""
        from openpyxl import Workbook

        bot, pw = _creds(ctx)
        wb = Workbook()
        names = [s for s in (sheets or []) if s.strip()] or ["Sheet1"]
        wb.active.title = names[0]
        for name in names[1:]:
            wb.create_sheet(name)
        await _put(bot, pw, path, _dump_xlsx(wb))
        return f"created {path}"

    @mcp.tool
    async def sheet_read(ctx: Context, path: str, sheet: str = "") -> list[list]:
        """Read a sheet's cells as a list of rows. Defaults to the first sheet."""
        bot, pw = _creds(ctx)
        wb = _load_xlsx(await _get(bot, pw, path))
        ws = _pick_sheet(wb, sheet)
        return [list(row) for row in ws.iter_rows(values_only=True)]

    @mcp.tool
    async def sheet_write_cell(ctx: Context, path: str, sheet: str, cell: str, value: str) -> str:
        """Set one cell (e.g. `B2`) on a sheet and save."""
        bot, pw = _creds(ctx)
        wb = _load_xlsx(await _get(bot, pw, path))
        ws = _pick_sheet(wb, sheet)
        try:
            ws[cell] = value
        except ValueError as exc:
            raise ToolError(f"invalid cell reference {cell!r}: {exc}") from exc
        await _put(bot, pw, path, _dump_xlsx(wb))
        return f"set {sheet}!{cell} in {path}"

    @mcp.tool
    async def sheet_append_row(ctx: Context, path: str, sheet: str, values: list) -> str:
        """Append a row of values to a sheet and save."""
        bot, pw = _creds(ctx)
        wb = _load_xlsx(await _get(bot, pw, path))
        ws = _pick_sheet(wb, sheet)
        ws.append(list(values))
        await _put(bot, pw, path, _dump_xlsx(wb))
        return f"appended a row to {sheet} in {path}"

    # --- documents (python-docx) ---

    @mcp.tool
    async def doc_create(ctx: Context, path: str, content: str = "") -> str:
        """Create a new `.docx` at `path`, optionally with initial text (one
        paragraph per line)."""
        bot, pw = _creds(ctx)
        await _put(bot, pw, path, _docx_bytes_from_text(content))
        return f"created {path}"

    @mcp.tool
    async def doc_read(ctx: Context, path: str) -> str:
        """Read a `.docx` as plain text (paragraphs joined by newlines)."""
        bot, pw = _creds(ctx)
        doc = _load_docx(await _get(bot, pw, path))
        return "\n".join(p.text for p in doc.paragraphs)

    @mcp.tool
    async def doc_write(ctx: Context, path: str, content: str) -> str:
        """Replace a `.docx`'s contents with `content` (one paragraph per line)."""
        bot, pw = _creds(ctx)
        await _put(bot, pw, path, _docx_bytes_from_text(content))
        return f"wrote {path}"

    @mcp.tool
    async def doc_append(ctx: Context, path: str, content: str) -> str:
        """Append paragraphs (one per line) to an existing `.docx`."""
        bot, pw = _creds(ctx)
        doc = _load_docx(await _get(bot, pw, path))
        for line in content.split("\n"):
            doc.add_paragraph(line)
        await _put(bot, pw, path, _dump_docx(doc))
        return f"appended to {path}"

    # --- listing + sharing ---

    @mcp.tool
    async def list_files(ctx: Context, path: str = "") -> list[str]:
        """List files/folders under `path` (default the bot's root)."""
        bot, pw = _creds(ctx)
        return await _run(nc.list_files, bot, pw, path)

    @mcp.tool
    async def create_share_link(
        ctx: Context,
        path: str,
        permission: str = "edit",
        expiry: str = "",
        password: str = "",
    ) -> str:
        """Create a public share link for a file the bot owns.

        `permission` is 'edit' or 'view'. Returns a `…/s/<token>` URL on the
        public host — the only place the public URL is used.
        """
        if permission not in ("edit", "view"):
            raise ToolError("permission must be 'edit' or 'view'")
        bot, pw = _creds(ctx)
        return await _run(
            nc.create_share_link, bot, pw, path,
            permission=permission,
            expiry=expiry or None,
            share_password=password or None,
        )

    return {
        "sheet_create": sheet_create,
        "sheet_read": sheet_read,
        "sheet_write_cell": sheet_write_cell,
        "sheet_append_row": sheet_append_row,
        "doc_create": doc_create,
        "doc_read": doc_read,
        "doc_write": doc_write,
        "doc_append": doc_append,
        "list_files": list_files,
        "create_share_link": create_share_link,
    }


# --- in-memory format helpers ---


def _dump_xlsx(wb) -> bytes:
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _load_xlsx(data: bytes):
    from openpyxl import load_workbook

    try:
        return load_workbook(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ToolError(f"not a readable .xlsx: {exc}") from exc


def _pick_sheet(wb, sheet: str):
    if not sheet:
        return wb.active
    if sheet not in wb.sheetnames:
        raise ToolError(f"sheet {sheet!r} not found (have: {', '.join(wb.sheetnames)})")
    return wb[sheet]


def _docx_bytes_from_text(content: str) -> bytes:
    """Build a fresh .docx from text (one paragraph per line) and serialize it.

    Shared by doc_create and doc_write — both start from an empty document and
    lay down the given lines. Empty content yields an empty document.
    """
    from docx import Document

    doc = Document()
    for line in content.split("\n") if content else []:
        doc.add_paragraph(line)
    return _dump_docx(doc)


def _load_docx(data: bytes):
    from docx import Document

    try:
        return Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ToolError(f"not a readable .docx: {exc}") from exc


def _dump_docx(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
